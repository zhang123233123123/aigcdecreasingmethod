import streamlit as st
import io
import requests
import tempfile
import os
from docx import Document
import base64
import re
import json
from stqdm import stqdm

# DeepSeek API 调用函数
def analyze_text_with_deepseek(text, api_key, ai_probability=50):
    """使用DeepSeek API分析文本的AI生成概率并提供优化建议，传入用户选择的AI率"""
    try:
        # 直接使用requests库调用API而不是OpenAI客户端
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # 根据用户选择的AI率构建不同的提示词
        system_prompt = f"""
        You are a professional text optimization assistant. Please analyze the following text and optimize it based on {ai_probability}% AI detection probability:

        Please optimize the text using these core strategies:
        1. Strategic Content Adjustment:
          - Identify paragraphs with obvious AI-generated characteristics and consider removing non-essential content
          - Rewrite or reduce high-risk sentences while maintaining core meaning
        
        2. Citation Strategy:
          - Add appropriate academic references and cross-references to improve text credibility
          - Use proper citation formats to reduce the probability of being identified as AI-generated
          - Rewrite high-risk content as quoted opinions from others
        
        3. Paragraph Restructuring:
          - Break higher-risk long paragraphs into multiple shorter paragraphs to avoid wholesale flagging
          - Logically divide large paragraphs into 2-3 smaller ones based on content flow
          - Add transition sentences between paragraphs to maintain logical coherence
        
        4. Word Choice and Sentence Variation:
          - Replace obvious AI-style vocabulary with more conversational expressions
          - Introduce natural language markers and conversational phrases (e.g., "actually," "to be honest," "it's worth noting")
          - Break up sentence pattern regularity with varied sentence lengths and thought patterns
          - Reconstruct parts of the discussion from a personal perspective (e.g., "I believe," "from this angle")

        Based on the text's AI probability of {ai_probability}%, please apply the following optimization strategy:
        """
        
        # 根据不同AI率范围添加具体优化策略
        if ai_probability > 70:
            system_prompt += """
            [DEEP OPTIMIZATION] Required changes:
              * Prioritize paragraph restructuring, breaking longer paragraphs into multiple shorter ones
              * Actively add academic citations, finding appropriate reference support for high-risk content
              * Thoroughly break sentence structure regularity, introducing natural variations in sentence length
              * Add conversational but academically acceptable phrases ("actually," "to be honest," "it's worth noting")
              * Moderately introduce transitions ("however," "nevertheless," "but looking back") and small-scale jump thinking
              * Reconstruct parts of the argument from a personal perspective ("I believe," "from this angle")
              * Replace overly textbook explanations with less "perfect" but still accurate explanations
              * Consider streamlining or physically removing non-core content
            """
        elif 60 <= ai_probability <= 70:
            system_prompt += """
            [MODERATE ADJUSTMENT] Required changes:
              * Apply paragraph restructuring to divide higher-risk paragraphs appropriately
              * Add 1-2 appropriate academic citations at key points
              * Preserve some original structure, but target overly regular paragraphs for adjustment
              * Add a small amount of personal insight and conversational connectors without affecting academic rigor
              * While maintaining the main arguments, slightly adjust the path of reasoning to make it more natural
              * Use synonym replacement to substitute obvious AI characteristic vocabulary with more humanized expressions
            """
        elif 50 <= ai_probability < 60:
            system_prompt += """
            [LIGHT OPTIMIZATION] Required changes:
              * Apply paragraph restructuring to individual sentences with obvious AI characteristics
              * Mainly preserve the original text, only fine-tuning the most obvious AI features
              * Replace 1-2 overly standardized expressions with more humanized tones
              * Adjust the structure of individual sentences while maintaining the overall appearance
              * Add appropriate conversational expressions like "actually," "it's worth noting," etc.
            """
        else:
            system_prompt += """
            [MAINTAIN ORIGINAL] AI probability below 50%:
              * The text already has good human writing characteristics
              * No need for extensive modifications; can maintain the original form
              * If needed, only adjust individual obviously mechanical expressions
              * Consider adding 1-2 personal opinion expressions
            """
        
        # 添加通用要求
        system_prompt += """
        All optimizations should:
        - Maintain basic academic writing standards, professionalism, and rigor
        - Ensure terminology accuracy remains unchanged
        - Maintain professional rigor while introducing conversational elements
        - Avoid excessive adjustments that distort content
        - Add appropriate personal perspective statements without compromising existing professionalism

        Please directly output the optimized text without explaining your modifications.
        """
        
        payload = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ]
        }
        
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers=headers,
            json=payload
        )
        
        # 检查响应状态
        if response.status_code == 200:
            response_data = response.json()
            optimized_text = response_data["choices"][0]["message"]["content"]
            return optimized_text
        else:
            st.error(f"API调用失败，状态码: {response.status_code}, 响应: {response.text}")
            return text
            
    except Exception as e:
        st.error(f"调用DeepSeek API时发生错误: {str(e)}")
        return text  # 出错时返回原文本

# 处理Word文档上传
def process_docx_upload(uploaded_file):
    """处理上传的Word文档，并保留原有格式"""
    try:
        # 创建临时文件保存上传的文档
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        
        # 读取文档
        doc = Document(tmp_path)
        os.unlink(tmp_path)  # 删除临时文件
        
        # 提取段落及其格式
        paragraphs = []
        styles = []  # 存储每个段落的样式信息

        for para in doc.paragraphs:
            if para.text.strip():  # 只处理非空段落
                paragraphs.append(para.text)
                
                # 保存段落样式信息
                para_style = {
                    'style_name': para.style.name,
                    'alignment': para.alignment,
                    'runs': []
                }
                
                # 保存每个run的格式信息
                for run in para.runs:
                    # 处理字体颜色
                    has_color = False
                    color_rgb = None
                    
                    # 安全地获取颜色信息
                    if hasattr(run, 'font') and hasattr(run.font, 'color'):
                        try:
                            font_color = run.font.color
                            if font_color and hasattr(font_color, 'rgb') and font_color.rgb:
                                has_color = True
                                # 获取颜色值，可能是整数或其他类型
                                color_rgb = font_color.rgb
                        except:
                            pass
                    
                    run_style = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size if run.font.size else None,
                        'font_name': run.font.name if run.font.name else None,
                        'has_color': has_color,
                        'color_rgb': color_rgb
                    }
                    para_style['runs'].append(run_style)
                
                styles.append(para_style)
        
        st.session_state.doc = doc
        st.session_state.paragraphs = paragraphs
        st.session_state.modified_paragraphs = paragraphs.copy()
        st.session_state.paragraph_styles = styles
        st.session_state.ai_probabilities = [50] * len(paragraphs)  # 默认概率值
        
        return True
    except Exception as e:
        st.error(f"处理文档时发生错误: {str(e)}")
        return False

# 获取文本的颜色标记
def get_color_class(ai_probability):
    """根据AI生成概率返回相应的颜色类名"""
    if ai_probability > 70:
        return "color-red"
    elif 60 <= ai_probability <= 70:
        return "color-orange"
    elif 50 <= ai_probability < 60:
        return "color-purple"
    else:
        return "color-black"

# 获取段落原始格式的HTML样式
def get_paragraph_style_html(para_idx):
    """根据保存的段落样式信息生成HTML样式"""
    if not hasattr(st.session_state, 'paragraph_styles') or para_idx >= len(st.session_state.paragraph_styles):
        return ""
    
    style = st.session_state.paragraph_styles[para_idx]
    style_str = ""
    
    # 添加对齐方式
    if style['alignment'] == 1:  # 居中
        style_str += "text-align: center; "
    elif style['alignment'] == 2:  # 右对齐
        style_str += "text-align: right; "
    elif style['alignment'] == 3:  # 两端对齐
        style_str += "text-align: justify; "
    
    # 如果段落有runs，使用第一个run的格式作为默认格式
    if style['runs'] and len(style['runs']) > 0:
        first_run = style['runs'][0]
        
        if first_run['bold']:
            style_str += "font-weight: bold; "
        if first_run['italic']:
            style_str += "font-style: italic; "
        if first_run['underline']:
            style_str += "text-decoration: underline; "
        if first_run['font_size'] is not None:
            # 将磅转换为像素（大致转换）
            pt_size = first_run['font_size'] / 12700  # 将EMU转换为磅
            px_size = pt_size * 1.33  # 磅转像素的大致转换
            style_str += f"font-size: {px_size}px; "
        if first_run['font_name'] is not None:
            style_str += f"font-family: '{first_run['font_name']}', sans-serif; "
        
        # 使用原始文本颜色，如果有
        if first_run['has_color'] and first_run['color_rgb']:
            try:
                # 简化颜色处理逻辑
                rgb = first_run['color_rgb']
                if isinstance(rgb, int):
                    # 如果是整数值，直接处理
                    r = rgb & 0xFF
                    g = (rgb >> 8) & 0xFF
                    b = (rgb >> 16) & 0xFF
                    style_str += f"color: rgb({r}, {g}, {b}); "
                elif hasattr(rgb, '_rgb'):
                    # 处理某些颜色对象
                    color_value = rgb._rgb
                    if color_value is not None:
                        style_str += f"color: #{color_value:06x}; "
                else:
                    # 默认使用黑色
                    style_str += "color: black; "
            except Exception as e:
                # 如果解析颜色出错，使用默认黑色
                style_str += "color: black; "
    
    return style_str

# 导出修改后的文档
def export_modified_doc():
    """导出修改后的Word文档，保留原格式"""
    try:
        if st.session_state.doc is None:
            st.error("没有可导出的文档")
            return None
        
        # 创建新文档
        doc = Document()
        
        # 获取修改后的文本段落
        modified_paragraphs = st.session_state.modified_paragraphs
        
        # 复制原文档的段落和格式
        for para_idx, para in enumerate(st.session_state.doc.paragraphs):
            if para.text.strip() and para_idx < len(modified_paragraphs):  # 只处理非空段落
                # 创建新段落并设置文本
                new_para = doc.add_paragraph()
                
                # 应用原段落的样式
                new_para.style = para.style
                
                # 判断是否有保存的样式信息
                if hasattr(st.session_state, 'paragraph_styles') and para_idx < len(st.session_state.paragraph_styles):
                    # 添加修改后的文本内容（不处理复杂格式，只添加纯文本）
                    new_para.add_run(modified_paragraphs[para_idx])
                else:
                    # 没有保存的样式信息，直接添加文本
                    new_para.add_run(modified_paragraphs[para_idx])
        
        # 创建内存中的字节流
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    except Exception as e:
        st.error(f"导出文档时发生错误: {str(e)}")
        return None

# 创建下载链接
def get_download_link(doc_bytes, filename="modified_document.docx"):
    """生成文档下载链接"""
    b64 = base64.b64encode(doc_bytes.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">点击下载修改后的文档</a>'
    return href

# 页面配置
st.set_page_config(
    page_title="降AIGC率",
    page_icon="✍️",
    layout="wide"
)

# 初始化会话状态
if 'input_text' not in st.session_state:
    st.session_state.input_text = ""
if 'output_text' not in st.session_state:
    st.session_state.output_text = ""
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""

# 自定义样式
st.markdown("""
<style>
    body {
        color: black;
        background-color: #121212;
    }
    .main {
        color: black !important;
    }
    .stTextArea textarea {
        min-height: 200px;
        background-color: white !important;
        border-radius: 10px !important;
        border: 1px solid #ddd !important;
        padding: 15px !important;
        color: black !important;
        font-size: 16px !important;
    }
    .text-card {
        background-color: white;
        border-radius: 10px;
        border: 1px solid #ddd;
        padding: 20px;
        margin: 10px 0;
        min-height: 200px;
        color: black !important;
    }
    .button-container {
        display: flex;
        gap: 10px;
        margin: 20px 0;
    }
    .stButton > button {
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        width: 100%;
    }
    .stButton > button:first-child {
        background-color: white;
        color: black;
        border: 1px solid #ddd;
    }
    .stButton > button:last-child {
        background-color: #4c6ef5;
        color: white;
        border: none;
    }
    .word-count {
        color: #999;
        font-size: 14px;
        margin-top: 8px;
    }
    .title {
        color: white;
        font-size: 32px;
        font-weight: bold;
        margin-bottom: 30px;
        text-align: center;
    }
    .section-title {
        color: white;
        font-size: 20px;
        font-weight: bold;
        margin-bottom: 15px;
    }
    .stAlert {
        background-color: #fff3cd;
        color: #856404;
        padding: 15px;
        border-radius: 10px;
        margin: 20px 0;
    }
    p {
        color: black !important;
    }
    .css-nahz7x {
        color: black !important;
    }
    div[data-testid="stMarkdownContainer"] > p {
        color: white !important;
    }
    .css-184tjsw p {
        color: white !important;
    }
    
    /* 显示字数的文本颜色 */
    div[data-testid="stMarkdownContainer"] > .word-count {
        color: #999 !important;
    }
</style>
""", unsafe_allow_html=True)

# 页面标题
st.markdown('<h1 class="title">降AIGC率</h1>', unsafe_allow_html=True)

# API密钥输入
api_key = st.text_input("请输入您的DeepSeek API密钥", 
                        value=st.session_state.api_key,
                        type="password",
                        help="需要DeepSeek API密钥才能分析和优化文本")

if api_key != st.session_state.api_key:
    st.session_state.api_key = api_key

# 左右两栏布局
col1, col2 = st.columns(2)

with col1:
    st.markdown('<div class="section-title">文档上传与编辑区</div>', unsafe_allow_html=True)
    
    # 文本输入区
    input_text = st.text_area("", 
                             value=st.session_state.input_text,
                             height=300,
                             placeholder="在此输入需要优化的文本...")
    
    if input_text != st.session_state.input_text:
        st.session_state.input_text = input_text
    
    # 显示字数
    word_count = len(input_text)
    st.markdown(f'<p class="word-count">{word_count}/1000 字符</p>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="section-title">文本对比与导出区</div>', unsafe_allow_html=True)
    
    # 显示优化后的文本
    st.markdown('<div class="text-card">', unsafe_allow_html=True)
    st.write(st.session_state.output_text if st.session_state.output_text else "优化后的文本将显示在这里...")
    st.markdown('</div>', unsafe_allow_html=True)

# 按钮区域
col1, col2 = st.columns([1, 1])

with col1:
    reset = st.button("重置", key="reset", use_container_width=True)
    if reset:
        st.session_state.output_text = ""
        st.experimental_rerun()

with col2:
    generate = st.button("一键生成", key="generate", use_container_width=True)
    if generate:
        if not st.session_state.api_key:
            st.error("请输入DeepSeek API密钥")
        elif not st.session_state.input_text:
            st.error("请输入需要优化的文本")
        else:
            with st.spinner("正在优化文本..."):
                optimized_text = analyze_text_with_deepseek(st.session_state.input_text, st.session_state.api_key)
                st.session_state.output_text = optimized_text
                st.experimental_rerun()

# 温馨提示
st.warning("为保护用户内容安全，段落处理的结果不会保存，请及时复制到自己的文件中。")
